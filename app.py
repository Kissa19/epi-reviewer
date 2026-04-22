import streamlit as st
import PyPDF2
import google.generativeai as genai
import io
from docx import Document
# ==========================================
# 1. ตั้งค่า System Prompt (กฎเกณฑ์ของ AI)
# ==========================================
SYSTEM_INSTRUCTION = """
บทบาท (Role):
คุณคือผู้เชี่ยวชาญด้านระบาดวิทยา และผู้ประเมินรายงานสอบสวนโรคระดับชาติ หน้าที่ของคุณคือการอ่าน "รายงานสอบสวนโรคฉบับสมบูรณ์ (Full Report)" และให้ข้อเสนอแนะเชิงวิจารณ์ (Constructive Feedback)

กฎเหล็กด้านความปลอดภัยของข้อมูล (CRITICAL RULE):
ห้ามมีข้อมูลส่วนบุคคล (PII): หากรายงานมีการระบุ ชื่อ, นามสกุล, เลขประจำตัวผู้ป่วย (HN), หรือ เลขบัตรประชาชน 13 หลัก ให้แจ้งเตือนเป็นข้อผิดพลาดร้ายแรง (Fatal Error) ทันที และแนะนำให้ผู้เขียนทำข้อมูลให้เป็นนิรนาม (Anonymization)

คำสั่ง (Task):
จงวิเคราะห์ข้อความรายงานสอบสวนโรคอย่างละเอียด โดยแบ่งการทำงานเป็น 2 ขั้นตอนหลัก:

ขั้นตอนที่ 1: ตรวจสอบความครบถ้วนของโครงสร้างรายงาน 14 หัวข้อหลัก
เช็คว่ามี: 1. ชื่อเรื่อง 2. ผู้รายงานและทีม 3. บทคัดย่อ 4. ความเป็นมา 5. วัตถุประสงค์ 6. วิธีการสอบสวน 7. ผลการสอบสวน 8. มาตรการควบคุม 9. วิจารณ์ผล 10. ปัญหา/ข้อจำกัด 11. ข้อเสนอแนะ 12. สรุปผล 13. กิตติกรรมประกาศ 14. เอกสารอ้างอิง

ขั้นตอนที่ 2: ประเมินคุณภาพและให้ข้อเสนอแนะรายหัวข้อ
* วัตถุประสงค์: 
    - สอบสวนเฉพาะราย: ต้องมี (เพื่อยืนยันการวินิจฉัย, ค้นหาแหล่งโรค/ลักษณะการเกิดโรค, หาแนวทางควบคุมป้องกัน)
    - สอบสวนการระบาด: ต้องมี (1.ยืนยันการวินิจฉัยและการระบาด 2.ศึกษาลักษณะทางระบาดวิทยาตาม Person, Time, Place 3.ค้นหาแหล่งโรค/วิธีการถ่ายทอด/ผู้สัมผัส 4.หามาตรการควบคุมป้องกันทั้งปัจจุบันและอนาคต)
* วิธีการสอบสวน: 
    - บุคคล (Person): ครอบคลุมผู้เกี่ยวข้อง ผู้สัมผัส หรือกลุ่มกิจกรรมร่วมกัน
    - สถานที่ (Place): ระบุสถานที่ที่เกี่ยวข้องสัมพันธ์กับปัจจัยบุคคล
    - เวลา (Time): การค้นหาผู้ป่วยเพิ่มเติมต้องครอบคลุมอย่างน้อย 1 เท่าของระยะฟักตัวยาวสุดก่อนเริ่มพบ Index case
    - อาการ/อาการแสดง: ต้องมีความเฉพาะเจาะจง (Specific) มากกว่าความถี่ที่พบบ่อย (Common) และอ้างอิงนิยามปี พ.ศ.2563 (สงสัย/เข้าข่าย/ยืนยัน)
* การศึกษาระบาดวิทยาเชิงวิเคราะห์ (Optional): เช็คการระบุรูปแบบการศึกษา (Cohort / Case-Control)
* ผลการสอบสวน: ต้องมีข้อมูลกระจายตามบุคคล, เวลา (Epidemic Curve ที่มี Time interval 1/3-1/8 ของระยะฟักตัว), สถานที่, ตาราง Bivariate/Multiple Logistic Regression, การสำรวจสิ่งแวดล้อม (ค่า อ.31, SI-2) และผลทางห้องปฏิบัติการ (Lab)
* มาตรการควบคุมป้องกันโรค: ต้องตอบวัตถุประสงค์ ระบุกิจกรรม (ใคร ทำอะไร ที่ไหน เมื่อไหร่ อย่างไร ทั้งระยะสั้นและยาว) และมีผลการติดตามเฝ้าระวังเชิงรุก
* วิจารณ์ผล: ห้ามเสนอตัวเลขซ้ำกับผลการสอบสวน ให้อธิบายความหมาย ความสอดคล้องของสาเหตุ/แหล่งโรค เปรียบเทียบกับเอกสารวิชาการ และระบุจุดอ่อน/ข้อจำกัด
* ข้อเสนอแนะ: ต้องสอดคล้องกับผลการสอบสวนและปัญหาที่พบ เพื่อการพัฒนาในอนาคต
* ปัญหาและข้อจำกัด: เขียนเป็นข้อๆ ถึงปัจจัยที่ทำให้การสอบสวนไม่เต็มตามวัตถุประสงค์
* สรุปผล: กระชับ ตอบทุกวัตถุประสงค์ ระบุโรค (สงสัย/เข้าข่าย/ยืนยัน) ปัจจัยเสี่ยง แหล่งโรค และระบุว่าควบคุมโรคได้หรือไม่ (ภายใน 2 เท่าระยะฟักตัวยาวสุดนับจากวันที่ลงสอบสวน)
รูปแบบผลลัพธ์ (Output Format): ให้แสดงผลเป็น Markdown แบ่งเป็น
1. คำแนะนำเบื้องต้น (แจ้งเตือน PII ถ้ามี)
2. 📋 ส่วนที่ 1: การตรวจสอบความครบถ้วน 14 หัวข้อหลัก (ใช้ ✅ หรือ ❌)
3. 📝 ส่วนที่ 2: ความคิดเห็นและข้อเสนอแนะรายหัวข้อ (ชื่นชมส่วนที่ดี และระบุข้อบกพร่อง)
4. 💡 ส่วนที่ 3: ตัวอย่างการปรับแก้
"""

# ==========================================
# 2. ฟังก์ชันสกัดข้อความ และส่งเข้า AI
# ==========================================
def extract_text_from_pdf(pdf_file):
    """สกัดข้อความจากไฟล์ PDF"""
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        if page.extract_text():
            text += page.extract_text() + "\n"
    return text

def analyze_report(api_key, text, report_type):
    """ส่งข้อมูลให้ Gemini วิเคราะห์ แบบป้องกัน Error 100%"""
    genai.configure(api_key=api_key)
    
    dynamic_prompt = SYSTEM_INSTRUCTION + f"\n\n**ข้อมูลเพิ่มเติมจากระบบ:** รายงานฉบับนี้เป็นการ {report_type}"
    prompt_text = f"โปรดประเมินรายงานสอบสวนโรคต่อไปนี้:\n\n{text}"
    
    try:
        # 1. ดึงรายชื่อโมเดลทั้งหมดที่ API Key ของคุณมีสิทธิ์ใช้จริงๆ
        available_models = []
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                # ตัดคำว่า models/ ออก เพื่อป้องกัน Error 404
                available_models.append(m.name.replace("models/", ""))
                
        if not available_models:
            return "❌ ข้อผิดพลาด: API Key ของคุณไม่มีสิทธิ์เข้าถึงโมเดลใดๆ เลย กรุณาสร้าง API Key ใหม่"

        # 2. คัดเลือกโมเดลที่ดีที่สุด (เรียงตามลำดับความฉลาด 1.5 Pro -> 1.5 Flash -> รุ่นอื่นๆ)
        target_model = available_models[0] # ตั้งค่าเริ่มต้นเป็นตัวแรกที่ระบบอนุญาต
        
        for name in available_models:
            if "gemini-1.5-pro" in name:
                target_model = name
                break
        else:
            for name in available_models:
                if "gemini-1.5-flash" in name:
                    target_model = name
                    break

        # 3. เริ่มทำการวิเคราะห์
        try:
            # พยายามรันแบบใช้ System Instruction (สำหรับ Gemini 1.5)
            model = genai.GenerativeModel(model_name=target_model, system_instruction=dynamic_prompt)
            response = model.generate_content(prompt_text)
            return response.text
            
        except Exception:
            # หากไลบรารีเก่าเกินไปจนไม่รู้จัก System Instruction ให้ใช้วิธีรวบข้อความแทน
            model = genai.GenerativeModel(model_name=target_model)
            full_prompt = f"คำสั่งของคุณคือ:\n{dynamic_prompt}\n\nข้อมูลรายงาน:\n{prompt_text}"
            response = model.generate_content(full_prompt)
            return response.text
            
    except Exception as e:
        return f"❌ พบข้อผิดพลาดของระบบ: {e}"
    
def create_word_doc(feedback_text):
    """ฟังก์ชันแปลงข้อความผลการประเมินเป็นไฟล์ Word"""
    doc = Document()
    doc.add_heading('ผลการประเมินรายงานสอบสวนโรค (Epi-Reviewer)', 0)
    
    # แยกข้อความตามบรรทัดเพื่อใส่ใน Word
    for line in feedback_text.split('\n'):
        doc.add_paragraph(line)
        
    # บันทึกไฟล์ลงในหน่วยความจำ (BytesIO) เพื่อเตรียมให้ดาวน์โหลด
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
# ==========================================
# 3. ส่วนหน้าตาแอปพลิเคชัน (UI)
# ==========================================
st.set_page_config(page_title="Epi-Reviewer | ระบบประเมินรายงาน", page_icon="📋", layout="wide")

st.title("📋 Epi-Reviewer: ผู้ช่วยตรวจประเมินรายงานสอบสวนโรค")
st.markdown("**โมดูลยกระดับงานวิชาการ (ส่วนหนึ่งของระบบ Epi-Analytic Pro)**")

# แถบด้านข้างสำหรับใส่ API Key
with st.sidebar:
    st.header("⚙️ ตั้งค่าระบบ")
    api_key_input = st.text_input("🔑 โปรดระบุ Gemini API Key ของคุณ", type="password", help="รับฟรีได้ที่ Google AI Studio")
    st.markdown("---")
    st.markdown("ระบบนี้ออกแบบตามคู่มือของ **กรมควบคุมโรค** เน้นการประเมิน 14 หัวข้อหลัก และหลักระบาดวิทยาเชิงลึก")

# พื้นที่หลัก แบ่งเป็น 2 คอลัมน์
col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("1. นำเข้าข้อมูลรายงาน")
    
    # ปุ่มเลือกประเภทการสอบสวน
    report_type = st.radio(
        "ประเภทการสอบสวนในรายงานฉบับนี้:",
        ["สอบสวนการระบาด (Outbreak)", "สอบสวนเฉพาะราย (Single Case)"],
        index=0
    )
    
    st.markdown("---")
    uploaded_file = st.file_uploader("2. อัปโหลดไฟล์รายงาน (PDF)", type=["pdf"])

with col2:
    st.subheader("3. ผลการประเมินจาก AI")
    
    if st.button("🚀 เริ่มตรวจสอบรายงาน", type="primary", use_container_width=True):
        if not api_key_input:
            st.warning("⚠️ กรุณาระบุ API Key ในแถบด้านข้างก่อนครับ")
        elif not uploaded_file:
            st.warning("⚠️ กรุณาอัปโหลดไฟล์รายงานฉบับสมบูรณ์ (PDF)")
        else:
            with st.spinner("⏳ ระบบกำลังอ่านเอกสารและวิเคราะห์โครงสร้าง 14 หัวข้อ... อาจใช้เวลาสักครู่"):
                try:
                    # 1. สกัดข้อความ
                    raw_text = extract_text_from_pdf(uploaded_file)
                    
                    # 2. วิเคราะห์
                    feedback = analyze_report(api_key_input, raw_text, report_type)
                    
                    # 3. แสดงผล
                    st.success("✅ วิเคราะห์เสร็จสมบูรณ์!")
                    
                    # ตกแต่งกล่องแสดงผล
                    with st.container(border=True):
                        st.markdown(feedback)
                        
                    # 4. สร้างปุ่มดาวน์โหลดเป็นไฟล์ Word
                    word_file = create_word_doc(feedback)
                    st.download_button(
                        label="💾 ดาวน์โหลดผลการประเมิน (Word)",
                        data=word_file,
                        file_name="Epi_Review_Feedback.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                        
                except Exception as e:
                    st.error(f"เกิดข้อผิดพลาดในการวิเคราะห์: {e}")
